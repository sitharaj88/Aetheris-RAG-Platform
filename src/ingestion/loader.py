"""
Multi-format document loader for the RAG1 system.

Supports PDF, DOCX, Markdown, plain text, and source code files.
Each loader returns a list of ``Document`` objects enriched with
metadata (source path, filename, file type, page numbers, etc.).
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from src.models import Document
from src.utils.logger import get_logger

log = get_logger(__name__)


class DocumentLoader:
    """
    Loads documents from individual files or entire directories.

    Usage::

        loader = DocumentLoader()
        docs = loader.load_file(Path("paper.pdf"))
        all_docs = loader.load_directory(Path("./corpus"), extensions=[".pdf", ".md"])
    """

    # Extensions treated as plain-text / source code
    TEXT_EXTENSIONS: set[str] = {
        ".txt", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".go", ".rs",
        ".html", ".css", ".json", ".yaml", ".yml", ".toml", ".sh", ".bat",
        ".rb", ".php", ".swift", ".kt",
    }

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def load_file(self, path: Path) -> list[Document]:
        """
        Load a single file and return one or more ``Document`` objects.

        Dispatches to the appropriate format-specific loader based on the
        file extension.
        """
        path = Path(path).resolve()
        if not path.exists():
            log.warning(f"File not found: {path}")
            return []

        ext = path.suffix.lower()
        log.info(f"Loading file: {path} (type={ext})")

        try:
            if ext == ".pdf":
                return self._load_pdf(path)
            elif ext == ".docx":
                return self._load_docx(path)
            elif ext == ".md":
                return self._load_markdown(path)
            elif ext in self.TEXT_EXTENSIONS:
                return self._load_text(path)
            else:
                log.warning(f"Unsupported extension '{ext}', attempting plain-text load")
                return self._load_text(path)
        except Exception as exc:
            log.error(f"Failed to load {path}: {exc}")
            return []

    def load_directory(
        self,
        directory: Path,
        extensions: Optional[list[str]] = None,
        recursive: bool = True,
    ) -> list[Document]:
        """
        Load all supported files under *directory*.

        Args:
            directory: Root directory to scan.
            extensions: Whitelist of extensions (e.g. ``[".pdf", ".md"]``).
                        If ``None``, all supported types are loaded.
            recursive: Walk subdirectories when ``True``.
        """
        directory = Path(directory).resolve()
        if not directory.is_dir():
            log.error(f"Not a directory: {directory}")
            return []

        pattern = "**/*" if recursive else "*"
        documents: list[Document] = []

        for file_path in sorted(directory.glob(pattern)):
            if not file_path.is_file():
                continue
            ext = file_path.suffix.lower()
            if extensions and ext not in extensions:
                continue
            docs = self.load_file(file_path)
            documents.extend(docs)

        log.info(f"Loaded {len(documents)} document(s) from {directory}")
        return documents

    # ------------------------------------------------------------------ #
    # Private format loaders
    # ------------------------------------------------------------------ #

    def _load_pdf(self, path: Path) -> list[Document]:
        """
        Extract text from each page of a PDF.

        Strategy:
        1. Try ``PyMuPDF (fitz)`` first — handles most PDFs well including
           designed PDFs, multi-column, etc.
        2. Fall back to ``pdfplumber`` if PyMuPDF not available.
        3. If still no text, attempt OCR fallback.
        """
        docs = []
        # Try PyMuPDF first (best for designed/complex PDFs)
        try:
            docs = self._load_pdf_pymupdf(path)
        except ImportError:
            log.debug("PyMuPDF not available, trying pdfplumber")
        except Exception as exc:
            log.warning(f"PyMuPDF failed for {path.name}: {exc}, trying pdfplumber")

        # Fallback to pdfplumber if no text extracted
        if not docs:
            try:
                docs = self._load_pdf_pdfplumber(path)
            except ImportError:
                log.debug("pdfplumber not available")
            except Exception as exc:
                log.error(f"pdfplumber failed for {path.name}: {exc}")

        # Fallback to OCR if still no text extracted
        if not docs:
            try:
                docs = self._load_pdf_ocr(path)
            except Exception as exc:
                log.error(f"OCR fallback failed for {path.name}: {exc}")

        return docs

    def _load_pdf_pymupdf(self, path: Path) -> list[Document]:
        """Extract text using PyMuPDF (fitz) — best for designed/complex PDFs."""
        import fitz  # PyMuPDF

        documents: list[Document] = []
        doc = fitz.open(str(path))
        total_pages = len(doc)

        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text("text")

            if not text.strip():
                # Try extracting with different method for image-heavy pages
                text = page.get_text("blocks")
                if isinstance(text, list):
                    text = "\n".join(
                        block[4] for block in text
                        if len(block) >= 5 and isinstance(block[4], str)
                    )

            if not text or not text.strip():
                continue

            documents.append(
                Document(
                    content=text.strip(),
                    source=str(path),
                    metadata={
                        "filename": path.name,
                        "file_type": "pdf",
                        "page_num": page_num + 1,
                        "total_pages": total_pages,
                        "source": str(path),
                    },
                )
            )

        doc.close()

        if documents:
            log.info(f"PDF {path.name}: {len(documents)} page(s) extracted via PyMuPDF")
        else:
            log.warning(f"PDF {path.name}: no text extracted by PyMuPDF")

        return documents

    def _load_pdf_pdfplumber(self, path: Path) -> list[Document]:
        """Fallback: extract text using pdfplumber."""
        import pdfplumber

        documents: list[Document] = []
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    continue
                documents.append(
                    Document(
                        content=text,
                        source=str(path),
                        metadata={
                            "filename": path.name,
                            "file_type": "pdf",
                            "page_num": page_num,
                            "total_pages": total_pages,
                            "source": str(path),
                        },
                    )
                )

        if documents:
            log.debug(f"PDF {path.name}: {len(documents)} page(s) via pdfplumber")

        return documents

    def _load_pdf_ocr(self, path: Path) -> list[Document]:
        """OCR fallback using rapidocr-onnxruntime."""
        try:
            from rapidocr_onnxruntime import RapidOCR
            import fitz
        except ImportError:
            log.error("rapidocr-onnxruntime or PyMuPDF (fitz) not installed. Cannot run OCR fallback.")
            return []

        log.info(f"Running OCR fallback for scanned PDF: {path.name}")
        try:
            engine = RapidOCR()
            doc = fitz.open(str(path))
            documents: list[Document] = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                # Render page to 150 DPI image bytes
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                
                # Perform OCR
                result, elapse = engine(img_data)
                if not result:
                    continue
                
                # Extract text lines and join them
                lines = [res[1] for res in result if res and len(res) >= 2]
                text = "\n".join(lines).strip()
                
                if text:
                    documents.append(
                        Document(
                            content=text,
                            source=str(path),
                            metadata={
                                "filename": path.name,
                                "file_type": "pdf",
                                "page_num": page_num + 1,
                                "total_pages": total_pages,
                                "source": str(path),
                                "extracted_via": "ocr",
                            },
                        )
                    )
            doc.close()
            return documents
        except Exception as e:
            log.error(f"OCR fallback failed for {path.name}: {e}")
            return []

    def _load_docx(self, path: Path) -> list[Document]:
        """Extract text from a DOCX file via python-docx."""
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [
            Document(
                content=full_text,
                source=str(path),
                metadata={
                    "filename": path.name,
                    "file_type": "docx",
                    "paragraph_count": len(paragraphs),
                    "source": str(path),
                },
            )
        ]

    def _load_markdown(self, path: Path) -> list[Document]:
        """Load a Markdown file as a single document."""
        text = path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            return []

        return [
            Document(
                content=text,
                source=str(path),
                metadata={
                    "filename": path.name,
                    "file_type": "markdown",
                    "source": str(path),
                },
            )
        ]

    def _load_text(self, path: Path) -> list[Document]:
        """Load a plain-text or source-code file."""
        text = path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            return []

        ext = path.suffix.lower().lstrip(".")
        return [
            Document(
                content=text,
                source=str(path),
                metadata={
                    "filename": path.name,
                    "file_type": ext if ext else "text",
                    "source": str(path),
                },
            )
        ]
